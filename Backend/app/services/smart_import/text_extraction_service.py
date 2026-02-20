# app/services/smart_import/text_extraction_service.py
"""
Text Extraction Service
Extraheert tekst uit PDF en DOCX bestanden voor Smart Import
TenderZen v3.0
"""
import io
import logging
import zipfile
from typing import Optional

logger = logging.getLogger(__name__)

# Probeer PyMuPDF te importeren (voor PDF)
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF not installed. Install with: pip install PyMuPDF")

# Probeer python-docx te importeren (voor DOCX)
try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")


class TextExtractionService:
    """
    Service voor het extraheren van tekst uit aanbestedingsdocumenten.
    
    Ondersteunde formaten:
    - PDF (via PyMuPDF/fitz)
    - DOCX (via python-docx)
    - ZIP (uitpakken en individuele bestanden verwerken)
    """
    
    async def extract(
        self,
        content: bytes,
        filename: str,
        mime_type: str = None
    ) -> str:
        """
        Extract tekst uit bestandsinhoud.
        
        Args:
            content: Raw bytes van het bestand
            filename: Originele bestandsnaam
            mime_type: MIME type van het bestand
            
        Returns:
            Geëxtraheerde tekst als string
        """
        extension = filename.lower().split('.')[-1]
        
        try:
            if extension == 'pdf' or mime_type == 'application/pdf':
                return await self._extract_pdf(content, filename)
            
            elif extension == 'docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_docx(content, filename)
            
            elif extension == 'zip' or mime_type in ['application/zip', 'application/x-zip-compressed']:
                return await self._extract_zip(content, filename)
            
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return f"[Bestandstype niet ondersteund: {extension}]"
                
        except Exception as e:
            logger.exception(f"Text extraction failed for {filename}: {e}")
            return f"[Fout bij extractie: {str(e)}]"
    
    async def _extract_pdf(self, content: bytes, filename: str) -> str:
        """Extract tekst uit PDF bestand."""
        if not HAS_PYMUPDF:
            return "[PDF extractie niet beschikbaar - installeer PyMuPDF: pip install PyMuPDF]"
        
        text_parts = []
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            logger.info(f"Extracting text from PDF: {filename} ({len(doc)} pages)")
            
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                
                if page_text.strip():
                    text_parts.append(f"\n--- Pagina {page_num} ---\n")
                    text_parts.append(page_text)
                
                # Probeer tabellen te extraheren
                try:
                    tables = page.find_tables()
                    if tables:
                        for table in tables:
                            table_data = table.extract()
                            if table_data:
                                text_parts.append("\n[Tabel]\n")
                                for row in table_data:
                                    row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                    text_parts.append(row_text + "\n")
                except:
                    pass  # Tabel extractie is best-effort
            
            doc.close()
            
            full_text = "".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {filename}")
            return full_text
            
        except Exception as e:
            logger.exception(f"PDF extraction error: {e}")
            raise ValueError(f"PDF extractie mislukt: {str(e)}")
    
    async def _extract_docx(self, content: bytes, filename: str) -> str:
        """Extract tekst uit DOCX bestand."""
        if not HAS_PYTHON_DOCX:
            return "[DOCX extractie niet beschikbaar - installeer python-docx: pip install python-docx]"
        
        text_parts = []
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            logger.info(f"Extracting text from DOCX: {filename}")
            
            # Extract paragrafen
            for para in doc.paragraphs:
                if para.text.strip():
                    if para.style and para.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {para.text}\n")
                    else:
                        text_parts.append(para.text + "\n")
            
            # Extract tabellen
            for table in doc.tables:
                text_parts.append("\n[Tabel]\n")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text + "\n")
                text_parts.append("\n")
            
            full_text = "".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {filename}")
            return full_text
            
        except Exception as e:
            logger.exception(f"DOCX extraction error: {e}")
            raise ValueError(f"DOCX extractie mislukt: {str(e)}")
    
    async def _extract_zip(self, content: bytes, filename: str) -> str:
        """Extract tekst uit ZIP bestand (recursief)."""
        text_parts = []
        
        try:
            with zipfile.ZipFile(io.BytesIO(content), 'r') as zip_file:
                logger.info(f"Extracting ZIP: {filename} ({len(zip_file.namelist())} files)")
                
                for inner_filename in zip_file.namelist():
                    if inner_filename.endswith('/') or inner_filename.startswith('__MACOSX'):
                        continue
                    
                    ext = inner_filename.lower().split('.')[-1]
                    if ext not in ['pdf', 'docx']:
                        continue
                    
                    try:
                        inner_content = zip_file.read(inner_filename)
                        inner_text = await self.extract(
                            content=inner_content,
                            filename=inner_filename,
                            mime_type=None
                        )
                        
                        if inner_text and not inner_text.startswith('['):
                            text_parts.append(f"\n\n{'='*40}\n")
                            text_parts.append(f"=== {inner_filename} ===\n")
                            text_parts.append(f"{'='*40}\n\n")
                            text_parts.append(inner_text)
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract {inner_filename}: {e}")
            
            if not text_parts:
                return "[ZIP bevat geen PDF of DOCX bestanden]"
            
            return "".join(text_parts)
            
        except Exception as e:
            logger.exception(f"ZIP extraction error: {e}")
            raise ValueError(f"ZIP extractie mislukt: {str(e)}")
    
    def check_dependencies(self) -> dict:
        """Check welke dependencies beschikbaar zijn."""
        return {
            'pdf': {'available': HAS_PYMUPDF, 'install': 'pip install PyMuPDF'},
            'docx': {'available': HAS_PYTHON_DOCX, 'install': 'pip install python-docx'},
            'zip': {'available': True, 'install': None}
        }
