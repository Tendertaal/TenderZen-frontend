# -*- coding: utf-8 -*-
"""
Markdown naar Tendertaal DOCX converter
TenderZen v3.2 - AI Document Export

Converteert gegenereerde markdown naar een professioneel .docx bestand
met Tendertaal huisstijl (paars #5B4B8A, rood #E01E37).

Ondersteunde markdown elementen:
  # H1  → Heading 1 (paars achtergrond, wit tekst)
  ## H2 → Heading 2 (paarse tekst, rode onderlijn)
  ### H3 → Heading 3 (rode tekst)
  **bold** → vetgedrukt
  - bullets → bulletlijst
  | tabel | → Word tabel met Tendertaal header
  Normale tekst → body paragraph
"""

import re
import io
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ── Tendertaal kleuren ──────────────────────────────────────────────
PAARS = RGBColor(0x5B, 0x4B, 0x8A)       # #5B4B8A
ROOD  = RGBColor(0xE0, 0x1E, 0x37)       # #E01E37
WIT   = RGBColor(0xFF, 0xFF, 0xFF)       # #FFFFFF
DONKER = RGBColor(0x1A, 0x1A, 0x2E)     # #1A1A2E
LICHT_PAARS = RGBColor(0xE8, 0xE5, 0xF0)  # #E8E5F0
GRIJS = RGBColor(0x64, 0x74, 0x8B)      # #64748B
LICHT_GRIJS = RGBColor(0xF5, 0xF5, 0xF5)  # #F5F5F5


def _hex_to_str(color: RGBColor) -> str:
    """RGBColor naar hex string voor XML."""
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_cell_bg(cell, hex_color: str):
    """Zet achtergrondkleur van een tabelcel via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, side: str, size: int = 4, color: str = 'CCCCCC'):
    """Zet celrand."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:color'), color)
    tcBorders.append(border)


def _add_bottom_border_to_para(para, color: str, size: int = 8):
    """Voeg onderlijn toe aan een paragraaf."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_shading_to_para(para, fill_color: str):
    """Voeg achtergrondkleur toe aan een paragraaf."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def _parse_inline(text: str) -> List[Tuple[str, bool]]:
    """
    Parseer inline markdown: **bold** en normale tekst.
    Geeft lijst van (tekst, is_bold) tuples terug.
    """
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    if not parts:
        parts = [(text, False)]
    return parts


def _add_run_with_inline(para, text: str, base_color: RGBColor = None,
                          base_size: int = 20, base_bold: bool = False):
    """Voeg tekst toe aan paragraaf met inline bold ondersteuning."""
    for segment, is_bold in _parse_inline(text):
        run = para.add_run(segment)
        run.font.name = 'Arial'
        run.font.size = Pt(base_size / 2)
        if base_color:
            run.font.color.rgb = base_color
        if is_bold or base_bold:
            run.bold = True


def _parse_table_line(line: str) -> List[str]:
    """Parseer een markdown tabelrij naar lijst van cellen."""
    line = line.strip().strip('|')
    cells = [c.strip() for c in line.split('|')]
    return cells


def _is_table_separator(line: str) -> bool:
    """Controleer of een regel een tabel-separator is (|---|---|)."""
    stripped = line.strip().strip('|')
    return bool(re.match(r'^[\s\-|:]+$', stripped))


def convert_markdown_to_docx(
    markdown: str,
    tender_naam: str = '',
    template_naam: str = '',
    opdrachtgever: str = '',
    gegenereerd_op: str = '',
) -> bytes:
    """
    Converteer markdown naar Tendertaal DOCX.
    Geeft bytes terug (klaar voor HTTP response of opslag).
    """
    doc = Document()

    # ── Pagina instellingen ────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)    # A4
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Standaard stijl ────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = DONKER

    # ── Header ────────────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    hpara = header.paragraphs[0]
    hpara.clear()
    hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hpara.add_run(f'Tendertaal  |  {template_naam}')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = PAARS
    run.bold = True
    if tender_naam:
        run2 = hpara.add_run(f'  |  {tender_naam}')
        run2.font.name = 'Arial'
        run2.font.size = Pt(8)
        run2.font.color.rgb = GRIJS
    _add_bottom_border_to_para(hpara, _hex_to_str(ROOD), size=6)

    # ── Footer ────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fpara = footer.paragraphs[0]
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = fpara.add_run(f'AI gegenereerd door TenderZen')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = GRIJS
    if gegenereerd_op:
        run2 = fpara.add_run(f'  |  {gegenereerd_op}')
        run2.font.name = 'Arial'
        run2.font.size = Pt(8)
        run2.font.color.rgb = GRIJS

    # ── Titelblok ─────────────────────────────────────────────────
    _add_title_block(doc, template_naam, tender_naam, opdrachtgever, gegenereerd_op)

    # ── Markdown parsen ───────────────────────────────────────────
    lines = markdown.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Tabel detectie
        if '|' in line and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            # Verzamel alle tabelregels
            table_lines = [line]
            i += 1  # sla separator over
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            _add_h1(doc, line[2:].strip())

        # Heading 2
        elif line.startswith('## ') and not line.startswith('### '):
            _add_h2(doc, line[3:].strip())

        # Heading 3
        elif line.startswith('### '):
            _add_h3(doc, line[4:].strip())

        # Bullet
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            _add_bullet(doc, text)

        # Genummerd item
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            _add_numbered(doc, text)

        # Horizontale lijn
        elif line.strip() in ('---', '***', '==='):
            _add_divider(doc)

        # Lege regel
        elif not line.strip():
            # Kleine spatie
            doc.add_paragraph('')

        # Normale tekst
        else:
            _add_body(doc, line)

        i += 1

    # ── Naar bytes ────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ── Bouw-functies ──────────────────────────────────────────────────

def _add_title_block(doc, template_naam, tender_naam, opdrachtgever, datum):
    """Voeg titelblok toe bovenaan het document."""
    # Paarse balk - titel
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _add_shading_to_para(p, _hex_to_str(PAARS))
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(template_naam or 'AI Gegenereerd Document')
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.color.rgb = WIT
    run.bold = True

    # Rode subtitel balk
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    _add_shading_to_para(p2, _hex_to_str(ROOD))
    p2.paragraph_format.left_indent = Cm(0.3)
    info_parts = []
    if tender_naam:
        info_parts.append(tender_naam)
    if opdrachtgever:
        info_parts.append(f'Opdrachtgever: {opdrachtgever}')
    if datum:
        info_parts.append(datum)
    run2 = p2.add_run('  |  '.join(info_parts) if info_parts else ' ')
    run2.font.name = 'Arial'
    run2.font.size = Pt(9)
    run2.font.color.rgb = WIT

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after  = Pt(6)


def _add_h1(doc, text: str):
    """Heading 1 - paarse achtergrond, witte tekst."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    _add_shading_to_para(p, _hex_to_str(PAARS))
    _add_run_with_inline(p, text, base_color=WIT, base_size=28, base_bold=True)


def _add_h2(doc, text: str):
    """Heading 2 - paarse tekst, rode onderlijn."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    _add_run_with_inline(p, text, base_color=PAARS, base_size=24, base_bold=True)
    _add_bottom_border_to_para(p, _hex_to_str(ROOD), size=6)


def _add_h3(doc, text: str):
    """Heading 3 - rode tekst."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    _add_run_with_inline(p, text, base_color=ROOD, base_size=22, base_bold=True)


def _add_body(doc, text: str):
    """Normale bodytekst met inline bold ondersteuning."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    _add_run_with_inline(p, text, base_color=DONKER, base_size=20)


def _add_bullet(doc, text: str):
    """Bulletpunt met paars bolletje."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    # Bullet karakter
    run_bullet = p.add_run('•  ')
    run_bullet.font.name = 'Arial'
    run_bullet.font.size = Pt(10)
    run_bullet.font.color.rgb = ROOD
    run_bullet.bold = True
    _add_run_with_inline(p, text, base_color=DONKER, base_size=20)


def _add_numbered(doc, text: str):
    """Genummerd item (simpele stijl, nummer al in text gestript)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.7)
    _add_run_with_inline(p, text, base_color=DONKER, base_size=20)


def _add_divider(doc):
    """Horizontale scheidingslijn."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _add_bottom_border_to_para(p, _hex_to_str(PAARS), size=4)


def _add_table(doc, table_lines: List[str]):
    """
    Voeg markdown tabel toe als Word tabel met Tendertaal opmaak.
    Eerste rij = header (paarse achtergrond, witte tekst).
    """
    if not table_lines:
        return

    # Parse header
    headers = _parse_table_line(table_lines[0])
    if not headers:
        return

    num_cols = len(headers)
    num_rows = len(table_lines)  # header al meegeteld

    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'

    # Kolombreedte berekenen (totaal ~16cm beschikbaar)
    total_width = int(16 * 914400 / 914400 * 914400)  # simplify
    col_width = Cm(16 / num_cols)

    # Header rij
    header_row = table.add_row()
    for ci, cell_text in enumerate(headers):
        cell = header_row.cells[ci]
        cell.width = col_width
        _set_cell_bg(cell, _hex_to_str(PAARS))
        for side in ['top', 'bottom', 'left', 'right']:
            _set_cell_border(cell, side, size=4, color=_hex_to_str(PAARS))
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(cell_text)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = WIT
        run.bold = True

    # Data rijen
    for row_idx, line in enumerate(table_lines):
        cells_data = _parse_table_line(line)
        # Zorg voor gelijk aantal kolommen
        while len(cells_data) < num_cols:
            cells_data.append('')

        row = table.add_row()
        bg = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'

        for ci in range(num_cols):
            cell = row.cells[ci]
            cell.width = col_width
            _set_cell_bg(cell, bg)
            for side in ['top', 'bottom', 'left', 'right']:
                _set_cell_border(cell, side, size=4, color='CCCCCC')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            cell_text = cells_data[ci] if ci < len(cells_data) else ''
            _add_run_with_inline(p, cell_text, base_color=DONKER, base_size=18)

    # Spacer na tabel
    doc.add_paragraph('')